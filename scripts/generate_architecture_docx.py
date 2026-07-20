from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "runtime-capacity-intelligence-architecture-summary.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(31, 41, 55)
MUTED = RGBColor(88, 99, 116)
HEADER_FILL = "F2F4F7"
CALLOUT_FILL = "F4F8FB"
BORDER = "DADFE8"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360, indent_dxa=120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2E74B5")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_field(paragraph, field_code: str) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    run._r.append(instr)

    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char)


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("Runtime Capacity Intelligence")
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    r = subtitle.add_run("Customer Validation Brief and Technical Summary")
    r.font.size = Pt(13)
    r.font.color.rgb = MUTED

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(14)
    r = meta.add_run(f"Customer-shareable project document | Updated {date.today().isoformat()} | Read-only UiPath Coded Web App")
    r.font.size = Pt(10)
    r.font.color.rgb = MUTED


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = DARK_BLUE
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)
    for run in p2.runs:
        run.font.color.rgb = INK
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, text in enumerate(headers):
        cell = hdr.cells[idx]
        set_cell_shading(cell, HEADER_FILL)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(text)
        r.bold = True
        r.font.color.rgb = INK

    for row_values in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row_values):
            cell = cells[idx]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.add_run(text)

    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)

    doc.add_paragraph()


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def add_footer(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Runtime Capacity Intelligence | Page ")
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    add_field(footer, "PAGE")


def build_doc() -> None:
    doc = Document()
    core = doc.core_properties
    core.author = "UiPath"
    core.last_modified_by = "UiPath"
    core.title = "Runtime Capacity Intelligence Customer Validation Brief"
    core.subject = "Read-only UiPath Coded Web App for unattended runtime capacity planning"
    core.keywords = "UiPath, Coded Web App, Orchestrator, Runtime Capacity, Customer Validation"
    core.comments = "Generated from maintained project documentation."

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    configure_styles(doc)
    add_footer(doc)
    add_title(doc)

    add_callout(
        doc,
        "Executive Summary",
        "Runtime Capacity Intelligence is a read-only UiPath Coded Web App for unattended robot runtime visibility, schedule contention analysis, machine template inventory, and what-if planning for new automations. The recommended next step is a controlled, read-only customer validation against a representative UiPath Cloud tenant.",
    )

    doc.add_heading("1. Current Working State", level=1)
    add_bullets(
        doc,
        [
            "React, TypeScript, Vite, Lucide React, and the UiPath TypeScript SDK are wired in the local app.",
            "The dashboard includes sign-in/sign-out, tenant/folder/machine-template filters, data health alerts, runtime heatmap day/week/month views, what-if planning, AI-assisted mock recommendations, schedule risk, and inventory.",
            "The app supports UiPath non-confidential External App browser sign-in, saved public connection profiles, read-only scope validation, and live connector diagnostics when tenant reads fail or return partial data.",
            "The live connector includes folder discovery and scoped reads for processes, jobs, machines, sessions, and triggers where customer permissions and endpoint scopes allow them.",
            "Demo data remains available for presentation, while live data should be validated against the target tenant before customer-facing findings are treated as authoritative.",
        ],
    )

    doc.add_heading("2. Customer Validation Request", level=1)
    p = doc.add_paragraph(
        "The pilot request is deliberately narrow: validate the app's read-only planning model against real UiPath Cloud operating patterns without changing the customer's automations."
    )
    p.paragraph_format.keep_with_next = True
    add_table(
        doc,
        ["Customer Input", "Purpose"],
        [
            ["UiPath Cloud tenant", "Use a representative environment with unattended schedules, machine templates, jobs, and runtime/session behavior."],
            ["External App client ID", "Enable browser OAuth through a customer-approved non-confidential External Application."],
            ["Redirect URI", "Match the exact local or deployed app URL used during validation."],
            ["Read-only scopes", "Permit approved folder, job, machine, robot/session, trigger, calendar, queue, license, or Data Fabric reads."],
            ["Pilot user assignment", "Scope visibility to customer-approved folders and machine templates through existing Orchestrator permissions."],
            ["Known risk windows", "Compare app heatmap and schedule-risk output against COE/admin expectations."],
        ],
        widths=[1.8, 4.6],
    )

    add_callout(
        doc,
        "Validation guardrail",
        "MVP behavior is advisory and read-only. The app does not start, stop, retry, kill, update, or reschedule jobs, and it does not store client secrets, UiPath credentials, tokens, raw queue payloads, or full robot logs.",
    )

    doc.add_heading("3. Architecture Overview", level=1)
    p = doc.add_paragraph(
        "The design separates UI composition from authentication, Orchestrator retrieval, DTO mapping, dashboard transforms, and planning logic. "
        "This keeps the UI readable and gives us one clear place to troubleshoot live data issues."
    )
    p.paragraph_format.keep_with_next = True

    add_table(
        doc,
        ["Layer", "Responsibility", "Primary Files"],
        [
            ["UI composition", "Page state, filters, panels, modals, and user interactions.", "src/App.tsx"],
            ["Configuration", "Scopes, timezone, date presets, storage keys, and API page limits.", "src/config/appConfig.ts"],
            ["Connection profiles", "Saved public connection details in browser local storage.", "src/lib/connectionProfiles.ts"],
            ["Authentication", "OAuth redirect flow, SDK lifecycle, sign-out, and token scope validation.", "src/lib/uipathAuth.ts"],
            ["Live probe facade", "Combines authentication and live retrieval into one UI-friendly result.", "src/lib/uipathLive.ts"],
            ["Orchestrator retrieval", "Reads folders, processes, jobs, machines, sessions, and triggers.", "src/lib/orchestratorRetriever.ts"],
            ["DTO mapping", "Normalizes SDK/OData responses into stable app contracts.", "src/lib/liveMappers.ts"],
            ["Dashboard transforms", "Builds live heatmap buckets, inventory rows, and schedule risks.", "src/lib/liveTransforms.ts"],
            ["Analytics and recommendations", "Utilization, risk grouping, peak demand, and rule-based recommendations.", "src/lib/analytics.ts; src/lib/recommendations.ts"],
            ["Type contracts", "Shared dashboard, connection, and live connector interfaces.", "src/types.ts; src/types/live.ts; src/types/connections.ts"],
        ],
        widths=[1.6, 2.7, 2.1],
    )

    doc.add_heading("4. Real Data Flow", level=1)
    add_numbered(
        doc,
        [
            "User clicks Sign in and selects or adds a saved UiPath connection profile.",
            "The app builds a ProbeConnectionConfig with platform URL, organization slug, tenant, client ID, redirect URI, and read-only scopes.",
            "uipathAuth.ts completes OAuth, initializes the UiPath SDK, and verifies required token scopes.",
            "orchestratorRetriever.ts reads folders first, then uses folder-scoped reads for processes, jobs, and triggers where possible.",
            "liveMappers.ts converts raw UiPath records into stable summaries.",
            "liveTransforms.ts converts those summaries into heatmap buckets, schedule risk records, and machine template inventory.",
            "App.tsx renders live data when available and surfaces connector warnings/errors before users trust the dashboard.",
        ],
    )

    doc.add_heading("5. Security and Data Handling", level=1)
    add_bullets(
        doc,
        [
            "Authentication uses a non-confidential UiPath External Application; no client secret belongs in the browser.",
            "Saved connection profiles contain only public configuration: platform URL, organization, tenant names, and client ID.",
            "The application is designed as read-only. Required scopes are centralized in src/config/appConfig.ts.",
            "OAuth tokens are handled by the UiPath TypeScript SDK and should never be logged, exported, stored in Data Fabric, or included in screenshots.",
            "For production history and performance, use Data Fabric/Data Service to store curated metrics rather than full raw job logs or queue payloads.",
        ],
    )

    doc.add_heading("6. Recommended Persisted Metric Model", level=1)
    add_table(
        doc,
        ["Entity", "Purpose"],
        [
            ["RuntimeCapacitySnapshot", "Hourly or 5-minute runtime used, available, total, and risk score."],
            ["MachineTemplateInventorySnapshot", "Configured and effective runtime capacity by machine template."],
            ["ProcessDurationMetric", "p50, p75, p90, and p95 duration by process, folder, and time bucket."],
            ["ScheduleRiskSnapshot", "Projected collision and SLA risk by folder, process, and schedule window."],
            ["WhatIfScenario", "Optional saved proposed automation scenarios and recommendation outcomes."],
            ["DataIngestionRun", "Refresh status, source counts, warnings, errors, and last successful load."],
            ["ApiErrorLog", "Sanitized API failures without tokens or sensitive payloads."],
        ],
        widths=[2.2, 4.2],
    )

    doc.add_heading("7. Live Data Troubleshooting Notes", level=1)
    add_callout(
        doc,
        "If sign-in succeeds but data is partial",
        "Start with the top Data Health band and Live Orchestrator Data preview, then inspect the connector messages from src/lib/orchestratorRetriever.ts. Most pilot findings should resolve to folder assignment, OAuth scope, endpoint support, retention, or pagination behavior before any UI change is considered.",
    )
    add_bullets(
        doc,
        [
            "Confirm the signed-in user has Orchestrator folder assignments in the selected tenant.",
            "Confirm the External App includes Folders, Execution/Jobs, Machines, Robots, and any approved Trigger/Schedule/Calendar/Queue/License read scopes needed by the target endpoints.",
            "Review probe messages from retrieveOrchestratorSnapshot before changing UI logic.",
            "If folder discovery fails but auth is valid, try alternate folder discovery endpoints or SDK methods in orchestratorRetriever.ts.",
            "Keep all connector errors visible in the data health band so business users understand whether they are seeing live, partial, or demo data.",
        ],
    )

    doc.add_heading("8. Deployment Path", level=1)
    add_table(
        doc,
        ["Step", "Command or Action"],
        [
            ["Install", "npm install"],
            ["Local development", "npm run dev"],
            ["Production build", "npm run build"],
            ["Pack Coded Web App", 'uip codedapp pack dist --name "Runtime Capacity Intelligence" --version 0.1.0'],
            ["Publish/deploy", "Use UiPath CLI codedapp publish/deploy after target app and folder are confirmed."],
        ],
        widths=[1.8, 4.6],
    )

    doc.add_heading("9. Customer Validation Success Criteria", level=1)
    add_bullets(
        doc,
        [
            "OAuth succeeds with the approved External App, redirect URI, organization, tenant, and read-only scopes.",
            "Folder discovery returns the expected pilot folders and clearly reports access gaps.",
            "Scoped Orchestrator reads return live records or actionable partial-access messages.",
            "Runtime heatmap and schedule risk outputs identify known high-contention windows.",
            "Machine template inventory aligns with customer-admin expectations for configured and effective capacity.",
            "What-if scenario output is understandable enough for release and scheduling discussion.",
            "The team can decide whether production should stay browser-only or add Data Fabric/Data Service metric persistence.",
        ],
    )

    doc.add_heading("10. Open Decisions", level=1)
    add_bullets(
        doc,
        [
            "Final read-only scope list for folders, triggers, licenses, queues, machines, robots, and jobs.",
            "Exact endpoint strategy for runtime/license utilization and machine-template capacity.",
            "Whether MVP stays browser-only or adds a backend/aggregation service.",
            "Whether Data Fabric/Data Service is sufficient for all target customer metric volumes.",
            "Whether heatmap buckets should be hourly, 15-minute, or tenant-configurable.",
            "Whether AI recommendations remain rules-only, use UiPath LLM Gateway, or use a hybrid approach.",
        ],
    )

    doc.add_heading("11. Source of Truth Files", level=1)
    add_bullets(
        doc,
        [
            "pdd.md - product and business requirements.",
            "sdd.md - solution design detail.",
            "docs/technical-architecture-and-deployment.md - expanded technical architecture and deployment notes.",
            "docs/code-architecture.md - current source-level file responsibility map.",
            "docs/uipath-runtime-concepts-and-api-reference.md - runtime concepts, endpoint inventory, field mapping, and customer validation questions.",
            "docs/uipath-runtime-concepts-and-api-reference.docx - customer-shareable Word version of the runtime/API reference.",
            "docs/runtime-capacity-intelligence-architecture-summary.docx - living Word summary for stakeholder sharing.",
        ],
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_doc()
