from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "uipath-runtime-concepts-and-api-reference.md"
OUTPUT = ROOT / "docs" / "uipath-runtime-concepts-and-api-reference.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(31, 41, 55)
MUTED = RGBColor(88, 99, 116)
HEADER_FILL = "E8EEF5"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360, indent_dxa=120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_field(paragraph, field_code: str) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    run._r.append(instr)
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Runtime Concepts and API Reference | Page ")
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    add_field(footer, "PAGE")

    core = doc.core_properties
    core.author = "UiPath"
    core.last_modified_by = "UiPath"
    core.title = "UiPath Runtime Concepts and API Reference"
    core.subject = "Runtime capacity, machine templates, host machines, serverless, and Orchestrator API fields"
    core.keywords = "UiPath, Orchestrator, Runtime, Machine Template, API Reference, Serverless"


def split_table_row(line: str) -> list[str]:
    cells = line.strip().strip("|").split("|")
    return [cell.strip().replace("<br>", "\n") for cell in cells]


def is_table_separator(line: str) -> bool:
    parts = split_table_row(line)
    return bool(parts) and all(re.fullmatch(r":?-{3,}:?", part.strip()) for part in parts)


def parse_inline(paragraph, text: str) -> None:
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = DARK_BLUE
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=1, cols=cols)
    table.style = "Table Grid"
    set_table_width(table)
    set_repeat_table_header(table.rows[0])

    for row_index, row_values in enumerate(rows):
        cells = table.rows[0].cells if row_index == 0 else table.add_row().cells
        for col_index in range(cols):
            cell = cells[col_index]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if row_index == 0:
                set_cell_shading(cell, HEADER_FILL)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if col_index < len(row_values):
                parse_inline(p, row_values[col_index])
                if row_index == 0:
                    for run in p.runs:
                        run.bold = True
                        run.font.color.rgb = INK

    doc.add_paragraph()


def flush_paragraph(doc: Document, lines: list[str]) -> None:
    if not lines:
        return
    p = doc.add_paragraph()
    parse_inline(p, " ".join(line.strip() for line in lines))


def flush_code(doc: Document, lines: list[str]) -> None:
    if not lines:
        return
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(45, 55, 72)


def build_doc() -> None:
    markdown = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_document(doc)

    paragraph_lines: list[str] = []
    table_rows: list[list[str]] = []
    code_lines: list[str] = []
    in_code = False
    title_written = False

    def flush_all() -> None:
        nonlocal paragraph_lines, table_rows, code_lines
        flush_paragraph(doc, paragraph_lines)
        paragraph_lines = []
        if table_rows:
            add_table(doc, table_rows)
            table_rows = []
        flush_code(doc, code_lines)
        code_lines = []

    for line in markdown:
        stripped = line.strip()
        if stripped.startswith("```"):
            if in_code:
                flush_code(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                flush_paragraph(doc, paragraph_lines)
                paragraph_lines = []
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            flush_paragraph(doc, paragraph_lines)
            paragraph_lines = []
            if is_table_separator(stripped):
                continue
            table_rows.append(split_table_row(stripped))
            continue

        if table_rows:
            add_table(doc, table_rows)
            table_rows = []

        if not stripped:
            flush_paragraph(doc, paragraph_lines)
            paragraph_lines = []
            continue

        heading = re.match(r"^(#{1,3})\s+(.*)$", stripped)
        if heading:
            flush_all()
            level = len(heading.group(1))
            text = heading.group(2)
            if level == 1 and not title_written:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(24)
                run.font.color.rgb = RGBColor(11, 37, 69)
                title_written = True
            else:
                doc.add_heading(text, level=min(level, 3))
            continue

        if stripped.startswith("- "):
            flush_paragraph(doc, paragraph_lines)
            paragraph_lines = []
            p = doc.add_paragraph(style="List Bullet")
            parse_inline(p, stripped[2:])
            continue

        if re.match(r"^\d+\.\s+", stripped):
            flush_paragraph(doc, paragraph_lines)
            paragraph_lines = []
            p = doc.add_paragraph(style="List Number")
            parse_inline(p, re.sub(r"^\d+\.\s+", "", stripped))
            continue

        paragraph_lines.append(stripped)

    flush_all()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_doc()

